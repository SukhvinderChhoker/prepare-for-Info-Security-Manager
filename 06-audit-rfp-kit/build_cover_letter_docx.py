"""RFP cover letter (.docx) - bracketed <fields> to fill at issuance."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from datetime import date

OUT = "RFP_Cover_Letter.docx"

doc = Document()
for section in doc.sections:
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

def H(text, size=12, bold=True, color=(15, 42, 74)):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold
    r.font.color.rgb = RGBColor(*color)
    r.font.name = "Calibri"

def N(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11); p.paragraph_format.space_after = Pt(4)

def B(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.size = Pt(11); r.bold = True

def BL(text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text); r.font.size = Pt(11)

# Letterhead
H("<NBFC Legal Name> Pvt. Ltd.", 16)
N("Regd Office: <Address>")
N("CIN: <U...>  GST: <...>  RBI Reg: <...>")

p = doc.add_paragraph(); r = p.add_run("\n"); p.add_run().font.size = Pt(11)
N(f"Date: {date.today().strftime('%d-%b-%Y')}")
N("To,")
B("<Firm Name>")
N("<Attn: Lead Partner / Country Head>")
N("<Address>")

N("Subject: Request for Proposal for <Engagement Title> (RFP ref: <RFP-No>)")

N("Dear <Lead Partner>,")

N("We, <NBFC Legal Name> Pvt Ltd, are pleased to invite your firm to submit a bid for the "
  "above-mentioned engagement. We are an RBI-registered non-banking financial company "
  "with operations in <geography>. We request that your firm carry out the scope defined "
  "in the appended RFP schedule, deliverables and timelines.")

N("This RFP package comprises:")
doc.add_paragraph("a) SOC 2 Auditor RFP - Schedule A", style="List Number").runs[0].font.size = Pt(11)
doc.add_paragraph("b) PCI DSS QSA RFP - Schedule B (if applicable)", style="List Number").runs[0].font.size = Pt(11)
doc.add_paragraph("c) Mandatory Bidder Disclosures - Schedule C", style="List Number").runs[0].font.size = Pt(11)
doc.add_paragraph("d) Negotiation tips and evaluation criteria - Schedule D", style="List Number").runs[0].font.size = Pt(11)

N("Submission deadline: <DD-MMM-YYYY, 17:00 IST>. Please upload your sealed response on "
  "the procurement portal <URL>, with hardcopy if requested.")

N("Your firm's proposal must include (over and above Schedule A/B):")
BL("Signed mutual NDA (Schedule E).")
BL("Independence attestation against Schedule C.")
BL("Sealed fee proposal (xlsx) on Schedule F template.")
BL("Audit methodology document - 10-15 pages.")
BL("Resumes of engagement partner, manager, in-charge consultant.")

N("Queries: All questions must be routed through the procurement portal no later than "
  "<DD-MMM-YYYY>. Answers will be published as FAQ on <DD-MMM-YYYY> without identifying "
  "the questioner.")

N("We look forward to receiving your proposal.")

N("Warm regards,")
N("<Name>")
N("<Title>, <NBFC>")
N("<email> | <phone>")

doc.save(OUT)
print("Wrote", OUT)
