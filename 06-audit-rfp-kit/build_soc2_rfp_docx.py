"""SOC 2 Auditor RFP template (.docx) for NBFCs."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "SOC2_Auditor_RFP_Template.docx"

doc = Document()

# Page setup A4
for section in doc.sections:
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)

def H(text, size=14, bold=True, color=(15, 42, 74)):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold
    r.font.color.rgb = RGBColor(*color)
    r.font.name = "Calibri"

def B(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10); r.bold = True

def N(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(2)

def BL(text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text); r.font.size = Pt(10)

def BLUE():
    p = doc.add_paragraph()
    r = p.add_run("____________________________________________________________")
    r.font.color.rgb = RGBColor(180, 180, 180)
    r.font.size = Pt(8)

# Cover
H("Request for Proposal - SOC 2 Type II Audit Services", 18)
N("Issuer: <NBFC Legal Name>      Issue date: <DD-MMM-YYYY>")
N("RFP ref: RFP-SOC2-<yyyy>-<seq>      Closing date: <DD-MMM-YYYY>")
BLUE()

# 1. Background
H("1. Background & scope")
N("<NBFC> is a Reserve Bank of India-registered non-banking financial company with <N> "
  "data principals and <N> employees. We process loan applications on AWS, integrate "
  "with a co-lending partner bank, and serve <N> downstream fintech customers via API. "
  "The SOC 2 Type II attestation is required for the customer-facing trust portal and "
  "fintech partner due-diligence.")
N("Scope period: <01-Jan-YYYY> to <31-Dec-YYYY>. Examination window preferred "
  "<Sep-YYYY> to <Feb-YYYY+1>.")

# 2. Standards coverage
H("2. Trust Services Criteria (TSC) coverage required")
BL("CC1 Control Environment")
BL("CC2 Communication & Information")
BL("CC3 Risk Assessment")
BL("CC4 Monitoring Activities")
BL("CC5 Control Activities")
BL("CC6 Logical and Physical Access")
BL("CC7 System Operations")
BL("CC8 Change Management")
BL("CC9 Risk Mitigation")
BL("A  Availability")
BL("C  Confidentiality")
BL("PI Processing Integrity (if processing loan calculations)")
BL("P  Privacy (if hosted customer PII in-scope)")

# 3. Auditor qualifications
H("3. Mandatory auditor qualifications")
BL("AICPA-licensed CPA firm with active peer review (must submit most recent peer review report).")
BL("SOC 2 practice lead with at least 5 attestations in last 36 months for cloud-native fintech / NBFCs.")
BL("Independence from <NBFC> and all material subservice organisations (must sign independence confirmation).")
BL("At least 2 partners / directors allocated to the engagement; named resumes required.")
BL("Carries professional indemnity cover >=Rs 50 Cr per claim for cyber-related errors/omissions.")
BL("Refs from 2 AICPA peer-reviewed NBFC / SOC 2 customers with comparable scope (will be contacted).")

# 4. Scope description
H("4. NBFC systems in-scope")
BL("Loan Origination System (LOS) on AWS EC2 / RDS, Mumbai (ap-south-1).")
BL("Mobile lending app - iOS + Android (>=1M installs).")
BL("Co-lending partner-bank API gateway.")
BL("Customer KYC storage with DPDP-tagged personal data.")
BL("Card disbursal flow via payment aggregator - check CDE overlap.")
BL("Subservice organisations: AWS (compute), <DRM Vendor> (DRM), <CDN> (web).")

# 5. Sampling & substantive testing
H("5. Sampling expectations")
N("Auditor's plan-of-test must commit to:")
BL("Sample sizes per CC family using AICPA AT-C §205 / §320 guidance.")
BL("Monetary-unit sampling for transaction-level controls (loan disbursal, repayment posting).")
BL("Full-population testing for change-management controls.")
BL("Sampling for SOC 2 bridge period covering <Oct-YYYY> to <Dec-YYYY>.")
N("Auditor will hand sampling tables to <NBFC> 30 days before fieldwork.")

# 6. Deliverables
H("6. Deliverables")
BL("Planning memo (Week 4).")
BL("Risk assessment matrix covering TSCs, services, and supporting ITGCs.")
BL("Internal-control walkthrough matrix (Week 8).")
BL("Draft Type II report with exceptions log (2 weeks before issuance).")
BL("Final SOC 2 Type II report (PDF + XBRL structured data if available).")
BL("Bridge letter covering <Jan-YYYY+1> to <date of final report>.")
BL("Subservice organisation carve-out (or inclusive) memo for AWS / DRM / CDN.")
BL("Management letter with advisory observations, severity-ranked.")

# 7. Engagement timeline
H("7. Engagement timeline")
N("RFP-issue: <date>        Q&A deadline: <date>")
N("Bid submission: <date>   Shortlist: <date>")
N("Bid presentation: <date> Chooser: <date>")
N("Planning kickoff: <date> Fieldwork start: <date>")
N("Draft report: <date>    Final report: <date>")

# 8. Fee structure
H("8. Fee structure")
N("Bid in INR. Provide:")
BL("Firm-fixed fee for SOC 2 Type II (full TSCs above).")
BL("Hourly rates for partner, director, manager, senior, staff (prorated).")
BL("Out-of-pocket ceiling (travel, copies).")
BL("Bridge-letter / SSAE-18 update fee per event.")
BL("Optional: SOC 2 + ISO 27001 + NIST CSF 2.0 integrated bid (preferred).")

# 9. Confidentiality & NDA
H("9. Confidentiality & NDA")
N("Auditor will sign mutual NDA before RFP submission. Bid is binding for 90 days.")

# 10. Insurance & indemnity
H("10. Insurance & indemnity")
BL("PI / E&O insurance >=Rs 50 Cr per claim with NBFC named additional insured.")
BL("Cyber insurance >=Rs 25 Cr.")
BL("No claim history breach certificate.")
BL("Indemnity language - auditor indemnifies NBFC for negligent findings in the report.")

# 11. Independence confirmations
H("11. Independence confirmations")
BL("No service, product, or sales relationship with <NBFC> for past 5 years.")
BL("No contingent fees.")
BL("Engagement partner has not served as NBFC's employee / director / advisor.")
BL("No internal-control design / implementation work for <NBFC> in scope year.")

# 12. Submission format
H("12. Required submission format")
BL("Cover letter (1 page, signed by Country Head).")
BL("Executive summary (2 pp) - engagement approach, key-staff, dates.")
BL("Resumes (engagement partner, manager, in-charge).")
BL("Plan-of-test (10 pp).")
BL("Fee proposal in sealed PDF + xlsx (Rate Card tab).")
BL("Recent peer review report (redacted).")
BL("2 NBFC / fintech SOC 2 client refs (will be phone-confirmed).")

# 13. Evaluation criteria
H("13. Evaluation criteria (weights shown)")
BL("Audit firm experience & NBFC track record - 25%")
BL("Quality of methodology & plan-of-test - 25%")
BL("Key-staff quality - 20%")
BL("Fee - 20%")
BL("Insurance / indemnity independence - 10%")

# 14. Contacts
H("14. RFP contacts")
N("RFP focal: <Name>, CFO Office - procurement@nbfc.in")
N("Technical: <Name>, CISO Office - ciso@nbfc.in")
N("Compliance: <Name>, DPO - dpo@nbfc.in")
N("All questions via procurement portal <URL> until <date>.")

BLUE()
N("<NBFC> reserves the right to reject any or all bids without assigning reasons. "
  "This RFP does not constitute an offer of engagement.")

doc.save(OUT)
print("Wrote", OUT)
