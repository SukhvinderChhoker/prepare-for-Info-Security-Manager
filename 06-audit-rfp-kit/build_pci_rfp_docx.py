"""PCI DSS QSA RFP template (.docx) for NBFCs with card / acquirer exposure."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor

OUT = "PCI_QSA_RFP_Template.docx"

doc = Document()
for section in doc.sections:
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)

def H(text, size=14, bold=True, color=(185, 28, 92)):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold
    r.font.color.rgb = RGBColor(*color)
    r.font.name = "Calibri"

def N(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10); p.paragraph_format.space_after = Pt(2)

def BL(text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text); r.font.size = Pt(10)

def B(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.size = Pt(10); r.bold = True

# Cover
H("Request for Proposal - PCI DSS v4.0 / v4.0.1 ROC Services", 18)
N("Issuer: <NBFC Legal Name>      Issue date: <DD-MMM-YYYY>")
N("RFP ref: RFP-PCI-<yyyy>-<seq>      Closing date: <DD-MMM-YYYY>")
B("Merchant Level: 1 (>= 6M Visa/MC transactions/year) OR Level 2 (1-6M) - confirm applicable level")

B("1. Scope of work")
N("Annual Report on Compliance (ROC) for the Cardholder Data Environment (CDE) defined below. "
  "Engagement must cover PCI DSS v4.0.1 mandatory requirements and the v3.2.1-defined future-dated "
  "controls that have already become effective.")
BL("PCI ROC (AoC) for full v4.0.1 mandatory and applicable customized-approach controls.")
BL("Quarterly ASV scan coordination (use approved PCI SSC ASV).")
BL("Annual penetration test scope review.")
BL("Annual internal-vulnerability scan review and sign-off.")
BL("Coordinated ROC + advisory for SAQ-eligible subservice organisations (payment gateway, DRM).")

B("2. CDE description")
BL("Tokenisation gateway from <Vendor> for debit-card disbursal.")
BL("Card vault - isolated VPC with HSM (FIPS 140-2 Level 3) holding PAN only, no CHD storage.")
BL("Reconciliation channel to <Acquirer> for settlement.")
BL("Card-flow endpoints on mobile app - PCI Scoped Mobile (Mobile v4).")
BL("Customer PAN never logged in plaintext - tokenised at SDK ingress only.")

B("3. Mandatory QSA qualifications")
BL("Active PCI SSC Qualified Security Assessor firm and named QSA employees on engagement.")
BL("NBFC or payment aggregator MC/VISA ROC track record - minimum 5 ROCs in last 36 months.")
BL("QSA-led onsite assessment (no remote-only for Level 1).")
BL("Two partner / lead QSA allocation; named resumes required.")
BL("PI / E&O insurance >=Rs 50 Cr; cyber insurance >=Rs 25 Cr.")
BL("Independence from all CDE vendors and the acquirer bank.")

B("4. v4.0.1 mandatory requirements - explicit attention")
BL("3.5.1 PAN encrypted when stored (AES-256 / TDES minimum).")
BL("4.1 strong cryptography for transmission over open networks (TLS 1.2+).")
BL("5.4.1 phishing protections - MFA on all CDE access, anti-phishing technical controls.")
BL("6.4.3 Targeted Risk Analysis (TRA) for each customised approach and PCI control with annual cadence.")
BL("8.4.2 MFA for all access into CDE - not just admins.")
BL("10.4.1.1 automated log review for CDE systems (daily / weekly).")
BL("11.6.1 change- and tamper-detection on payment pages.")
BL("12.3.1 inventory of all CDE components and sensitive data flows.")

B("5. Sampling & substantive testing")
BL("TRA methodology must be shared at bid stage; customised-approach controls require lead-QSA sign-off per control.")
BL("RBA (Risk-Based Approach) scoping - bidder to declare which eligible controls are RBA-ed and provide residual-risk rationale.")
BL("Sample sizes for transaction-level controls (issue, settlement, refund) - minimum 25 each.")
BL("Penetration test review to align with ROC period (must be no older than 6 months at ROC-issuance).")

B("6. Deliverables")
BL("Pre-engagement scoping memo (Week 3).")
BL("CDE network diagram red-lined review (Week 6).")
BL("Pre-attestation gap memo (Week 12).")
BL("Draft ROC with full v4 workbook (Week 18).")
BL("Final ROC + AoC signed by lead QSA.")
BL("Quarterly ASV scan summary (4 events/year).")
BL("Pen test review memo (signed).")
BL("Customised Approach Workbook for any controls not met by defined approach.")

B("7. Engagement timeline")
B("RFP-issue: <date>        Q&A deadline: <date>")
B("Bid submission: <date>   Shortlist: <date>")
B("Bid presentation: <date> Chooser: <date>")
B("Scoping kickoff: <date> ASV scan Qx:")
B("Fieldwork: <date> to <date>       ROC final: <date>")

B("8. Fee structure (INR)")
BL("Fixed-fee ROC including all v4 + AC workbooks + AoC.")
BL("Hourly rates: Lead QSA, QSA, Senior, Staff.")
BL("Travel / OOP ceiling.")
BL("Optional: Integrated QSA + SSAE18 (SOC 2) bid - preferred to reduce subservice duplication.")
BL("Optional: Pen-test review and ASV subscription fee.")

B("9. Compliance & escalations")
BL("All non-compliance findings rated 'high' or 'critical' must be reported to CISO within 24h of detection.")
BL("Breach-imminent findings must trigger acquirer 24h notification clock - bidder must confirm they will draft the acquirer template.")
BL("Bidder must align with our CERT-IN 6h and DPDP 72h breach-clocks. Acquirer 24h sits between CERT-IN 6h and DPDP 72h - bidder's draft must fit.")

B("10. Independence confirmations")
BL("No CDE vendor relationship in past 5 years.")
BL("No acquirer-bank relationship affecting this ROC.")
BL("Lead QSA has not been NBFC employee, director, or contractor in past 5 years.")
BL("No marketing / referral payments between QSA and card networks for this engagement.")

B("11. Submission format")
BL("Cover letter (1 page) signed by Country Manager.")
BL("Executive summary - approach, lead QSA bios, schedule, fee.")
BL("Resumes for QSA-led team.")
BL("Plan-of-test (15 pp) including TRA methodology and RBA rationale.")
BL("Sealed fee proposal (xlsx Rate Card).")
BL("Recent 3 ROCs (redacted) for NBFC / payment-aggregator clients.")
BL("Insurance certificates.")
BL("Sample of customised-approach workbook from a previous ROC.")

B("12. Evaluation criteria")
BL("QSA firm NBFC / payment track record - 25%")
BL("Quality of plan-of-test & TRA methodology - 25%")
BL("Quality of QSA team - 20%")
BL("Fee - 20%")
BL("Insurance / indemnity / independence - 10%")

B("13. Contacts")
N("RFP focal: <Name>, CFO Procurement - procurement@nbfc.in")
N("PCI focal: <Name>, CISO - ciso@nbfc.in")
N("Compliance: <Name>, DPO - dpo@nbfc.in")
N("All queries via procurement portal <URL> until <date>.")

N("<NBFC> reserves the right to reject any or all bids without assigning reasons. "
  "This RFP does not constitute an offer of engagement.")

doc.save(OUT)
print("Wrote", OUT)
