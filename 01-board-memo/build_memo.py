"""
NBFC CISO Board Memo - 1-page executive summary (landscape A4 for safe fit).
Builds PDF + DOCX.
"""
from reportlab.lib.pagesizes import landscape, A4
from reportlab.lib.styles import ParagraphStyle, StyleSheet1
from reportlab.lib.units import cm
from reportlab.lib.colors import HexColor, white
from reportlab.lib.enums import TA_LEFT
from reportlab.platypus import Paragraph, Spacer, Table, TableStyle, KeepTogether
from reportlab.platypus.frames import Frame
from reportlab.platypus.doctemplate import PageTemplate, BaseDocTemplate
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INK     = HexColor("#1E293B")
PRIMARY = HexColor("#0F766E")
ACCENT  = HexColor("#A21CAF")
SUBTLE  = HexColor("#475569")
LIGHT   = HexColor("#F0FDFA")
RULE    = HexColor("#CCFBF1")

DATE_LINE = datetime.now().strftime("%d %B %Y")

META = [
    ("TO:",      "Risk Committee / Audit Committee of the Board"),
    ("FROM:",    "Chief Information Security Officer (CISO)"),
    ("SUBJECT:", "Joint SOC 2 + PCI DSS v4.0 Compliance Programme"),
    ("DATE:",    DATE_LINE),
    ("REF:",     "NBFC/2026/CISO/JOINT-PROG/001"),
]

SECTIONS = [
    ("Purpose",
     "Seek approval to run SOC 2 Type II and PCI DSS Service Provider Level 1 ROC as a single "
     "integrated programme. Combined run delivers ~30% lower engineering-hours, ~15-20% lower "
     "audit fees, and a unified board reporting cadence. NBFC touches card data AND enterprise "
     "customers; running separately doubles work and splits the audit month in two."),
    ("Recommendation",
     "Approve the joint programme. Authorise a single SteerCo (CISO + CRO + CFO + CTO), single "
     "GRC (Vanta or Drata), single 18-month roadmap, and budget of INR 1.9-3.1 crore for cycle 1. "
     "Authorise the CISO to launch by week 4 with vendor selection complete by week 6."),
    ("Why now",
     "1) Enterprise customers (banks, B2B SaaS, marketplaces) demand SOC 2 Type II before "
     "onboarding - current absence blocks sales. 2) Acquirer / card networks mandate annual PCI "
     "ROC for NBFC card data. 3) RBI IT Framework 2023 + DPDP Act 2023 + Tokenisation 2021/2023 "
     "layer regulatory weight. Two parallel programmes double engineering load; one halves it."),
]

ACTIONS = [
    ("Approve",     "Joint SOC 2 + PCI DSS programme structure + 18-month roadmap",   "Today",   "Risk Committee"),
    ("Authorise",   "INR 1.9-3.1 crore total cycle 1 budget",                            "Today",   "Risk Committee"),
    ("Charter",     "Combined SteerCo chaired by CRO; CISO leads delivery",              "Week 2",  "CISO + CRO"),
    ("Vendor",      "Select QSA + CPA auditor pair + GRC tool",                           "Week 4",  "CFO + CISO"),
    ("Programme",   "Mobilise SteerCo + evidence repository + RACI",                       "Week 6",  "CISO"),
    ("Window",      "Begin first joint observation window",                                "Week 12", "CISO"),
]

KPIS = [
    ["KPI", "Separate run", "Joint run", "Delta"],
    ["Engineering FTE-hours / qtr", "1,400-1,800", "950-1,200", "-30%"],
    ["Annual audit fees",            "INR 1.8-3.0 cr", "INR 1.5-2.4 cr", "-15 to -20%"],
    ["Audit findings at first cycle","12-15",         "7-9",          "-40%"],
    ["SOC 2 + PCI delivery days",    "270 days",      "210 days",     "-22%"],
    ["Enterprise sales cycle",       "120 days",      "75 days",      "-37%"],
    ["Compliance dashboards BAU",    "5-7",           "1 unified cockpit","Single source"],
]

RISKS = [
    ["Risk", "Description", "Mitigation"],
    ["Two parallel programmes", "Engineering burnout + duplicated controls",
     "Joint SteerCo + shared evidence + shared GRC tool"],
    ["Audit date collision",    "SOC 2 + PCI in same month",
     "Sequential fieldwork; PCI Sep-Oct, SOC 2 Jan-Feb"],
    ["Vendor AOC expiry slip",  "Vendor SOC 2 or PCI AOC expires quietly",
     "Quarterly CISO review + automated calendar alerts"],
    ["RBI 2-hour timer",        "Cyber incident hits the cliff",
     "Quarterly tabletop + parallel clocks playbook"],
]

DECISIONS = [
    ("Approve",   "Joint SOC 2 + PCI DSS programme structure and 18-month roadmap"),
    ("Authorise", "INR 1.9-3.1 crore cycle 1 budget"),
    ("Charter",   "Combined SteerCo chaired by CRO, CISO-led execution"),
    ("Schedule",  "First joint observation window start by 30 Sep 2026"),
    ("Authorise", "CISO to negotiate single GRC tool capex <= INR 65 lakh"),
]

CLOSER = "Programme charter with detailed milestones, vendor shortlist and contingencies will be "\
         "circulated post-approval under separate cover. Submitted for Risk Committee review."

SIGNOFF = "- CISO  |  Reviewed: CFO  |  Reviewed: CRO"

# ===========================================================================
# PDF - landscape A4
# ===========================================================================
OUT_PDF = "NBFC_CISO_Board_Memo_Compliance_Programme.pdf"
PAGE = landscape(A4)  # 29.7 x 21.0 cm
PAGE_W, PAGE_H = PAGE

HEAD = ParagraphStyle('Head', fontName='Helvetica-Bold', fontSize=14, leading=16,
                     textColor=PRIMARY, spaceBefore=0, spaceAfter=4, alignment=TA_LEFT)
H    = ParagraphStyle('H',    fontName='Helvetica-Bold', fontSize=10, leading=12,
                      textColor=PRIMARY, spaceBefore=0, spaceAfter=2, alignment=TA_LEFT)
Body = ParagraphStyle('Body', fontName='Helvetica',       fontSize=8, leading=10,
                      textColor=INK, alignment=TA_LEFT, spaceAfter=2)
Tag  = ParagraphStyle('Tag',  fontName='Helvetica-Bold',  fontSize=7, leading=9,
                      textColor=white, alignment=TA_LEFT)

def head_meta():
    rows = [[Paragraph(f"<b>{k}</b>", Body), Paragraph(v, Body)] for k,v in META]
    t = Table(rows, colWidths=[2.2*cm, 24.5*cm])
    t.setStyle(TableStyle([
        ("VALIGN",(0,0),(-1,-1),"TOP"),
        ("LEFTPADDING",(0,0),(-1,-1), 0),
        ("RIGHTPADDING",(0,0),(-1,-1), 0),
        ("TOPPADDING",(0,0),(-1,-1), 0),
        ("BOTTOMPADDING",(0,0),(-1,-1), 0),
    ]))
    return t

def section_table(title, body):
    head = Paragraph(f"<b>{title}</b>", H)
    body_p = Paragraph(body.replace("\n","<br/>"), Body)
    inner = Table([[head, body_p]], colWidths=[3.5*cm, 23.2*cm])
    inner.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(0,0),LIGHT),
        ("LEFTPADDING",(0,0),(-1,-1), 4),
        ("RIGHTPADDING",(0,0),(-1,-1), 4),
        ("TOPPADDING",(0,0),(-1,-1), 2),
        ("BOTTOMPADDING",(0,0),(-1,-1), 2),
        ("LINEBEFORE",(0,0),(-1,-1), 1.5, PRIMARY),
        ("VALIGN",(0,0),(-1,-1),"TOP"),
    ]))
    return inner

def tag_para(text):
    return Paragraph(f"<font color='white'><b>{text}</b></font>", Tag)

def make_table(rows, widths):
    t = Table(rows, colWidths=[w*cm for w in widths])
    t.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0), PRIMARY),
        ("TEXTCOLOR",(0,0),(-1,0), white),
        ("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),
        ("FONTSIZE",(0,0),(-1,0), 7.6),
        ("FONTSIZE",(0,1),(-1,-1), 7.2),
        ("ALIGN", (0,0),(-1,0),"LEFT"),
        ("ROWBACKGROUNDS",(0,1),(-1,-1), [white, LIGHT]),
        ("GRID",(0,0),(-1,-1), 0.25, RULE),
        ("VALIGN",(0,0),(-1,-1),"TOP"),
        ("LEFTPADDING",(0,0),(-1,-1), 3),
        ("RIGHTPADDING",(0,0),(-1,-1), 3),
        ("TOPPADDING",(0,0),(-1,-1), 1),
        ("BOTTOMPADDING",(0,0),(-1,-1), 1),
    ]))
    return t

def wrap_table(title, t):
    return KeepTogether([Paragraph(f"<b>{title}</b>", H), t, Spacer(1, 0.08*cm)])

story = []
story.append(Paragraph("Board Memo - Joint SOC 2 + PCI DSS Compliance Programme", HEAD))
story.append(head_meta())
story.append(Spacer(1, 0.05*cm))
for t_,b in SECTIONS:
    story.append(section_table(t_, b))
    story.append(Spacer(1, 0.04*cm))
story.append(wrap_table("Quantified wins (joint vs separate)", make_table(KPIS, [7.0, 6.0, 6.0, 7.7])))
story.append(wrap_table("Risk Committee ask - six decisions",
                        make_table([["Action","Detail","When","Owner"]] + list(ACTIONS),
                                   [2.6, 12.8, 3.5, 7.8])))
story.append(wrap_table("Top four risks and mitigations",
                        make_table(RISKS, [5.5, 11.5, 9.7])))
story.append(wrap_table("Specific decisions requested",
                        make_table([["Action verb","Decision text"]] + list(DECISIONS),
                                   [3.0, 23.7])))
story.append(Spacer(1, 0.05*cm))
story.append(Paragraph("<b>Closing:</b> " + CLOSER, Body))
story.append(Spacer(1, 0.03*cm))
story.append(Paragraph(SIGNOFF, Body))

doc = BaseDocTemplate(
    OUT_PDF, pagesize=PAGE,
    leftMargin=0.6*cm, rightMargin=0.6*cm,
    topMargin=0.5*cm, bottomMargin=0.5*cm,
    title="NBFC Board Memo - Joint SOC 2 + PCI DSS",
    author="Cyber GRC Research"
)
def _on_page(canvas, doc_):
    canvas.saveState()
    canvas.setStrokeColor(RULE); canvas.setLineWidth(0.5)
    canvas.rect(0.5*cm, 0.5*cm, PAGE_W-1.0*cm, PAGE_H-1.0*cm)
    canvas.setFont("Helvetica-Oblique", 7)
    canvas.setFillColor(SUBTLE)
    canvas.drawString(1.0*cm, 0.25*cm, f"{META[4][1]}  -  page {doc_.page}")
    canvas.drawRightString(PAGE_W-1.0*cm, 0.25*cm, "Landscape A4 - safe 1-page fit")
    canvas.restoreState()
frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="m")
doc.addPageTemplates([PageTemplate(id="m", frames=[frame], onPage=_on_page)])
doc.build(story)
import re
pages = len(re.findall(rb'/Type\s*/Page[^s]', open(OUT_PDF,'rb').read()))
print(f"Wrote {OUT_PDF}  ({pages} page)")

# ===========================================================================
# DOCX (also widened for landscape printable layout)
# ===========================================================================
OUT_DOCX = "NBFC_CISO_Board_Memo_Compliance_Programme.docx"
d = Document()
s = d.sections[0]
s.orientation = 1  # landscape
s.page_height, s.page_width = Cm(21.0), Cm(29.7)
s.top_margin = Cm(0.8); s.bottom_margin = Cm(0.8)
s.left_margin = Cm(1.0); s.right_margin = Cm(1.0)

t = d.add_paragraph()
t.paragraph_format.space_after = Pt(2)
run = t.add_run("Board Memo - Joint SOC 2 + PCI DSS Compliance Programme")
run.font.size = Pt(14); run.font.bold = True
run.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)

mt = d.add_table(rows=5, cols=2)
mt.autofit = False
mt.columns[0].width = Cm(2.5)
mt.columns[1].width = Cm(24.5)
for i,(k,v) in enumerate(META):
    c1 = mt.rows[i].cells[0]; c2 = mt.rows[i].cells[1]
    c1.width = Cm(2.5);     c2.width = Cm(24.5)
    p1 = c1.paragraphs[0]
    r1 = p1.add_run(k); r1.font.size = Pt(9); r1.font.bold = True
    r1.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)
    p2 = c2.paragraphs[0]
    r2 = p2.add_run(v); r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

def add_section_p(title, body):
    h = d.add_paragraph()
    h.paragraph_format.space_before = Pt(4); h.paragraph_format.space_after = Pt(1)
    r = h.add_run(title); r.font.size = Pt(10); r.font.bold = True
    r.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)
    for line in body.split("\n"):
        p = d.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(line); run.font.size = Pt(10)

for t_,b in SECTIONS:
    add_section_p(t_, b)

def stylize_header(cell, text):
    cell.paragraphs[0].clear()
    r = cell.paragraphs[0].add_run(text)
    r.font.size = Pt(8.5); r.font.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), "0F766E"); tcPr.append(shd)
def stylize_cell(cell, text):
    cell.paragraphs[0].clear()
    r = cell.paragraphs[0].add_run(text)
    r.font.size = Pt(8.5); r.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

def add_table(rows, widths_cm):
    tbl = d.add_table(rows=len(rows), cols=len(widths_cm))
    tbl.style = "Light Grid Accent 1"
    for ci,w in enumerate(widths_cm):
        for c in tbl.columns[ci].cells: c.width = Cm(w)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri].cells[ci]
            if ri == 0: stylize_header(cell, val)
            else: stylize_cell(cell, val)

add_section_p("Quantified wins (joint vs separate)", "")
add_table(KPIS, [7.0, 6.0, 6.0, 7.7])

add_section_p("Risk Committee ask - six decisions", "")
add_table([["Action","Detail","When","Owner"]] + list(ACTIONS), [2.6, 12.8, 3.5, 7.8])

add_section_p("Top four risks and mitigations", "")
add_table(RISKS, [5.5, 11.5, 9.7])

add_section_p("Specific decisions requested", "")
add_table([["Action verb","Decision text"]] + list(DECISIONS), [3.0, 23.7])

add_section_p("Closing", "")
p = d.add_paragraph()
r = p.add_run(CLOSER); r.font.size = Pt(10)

p = d.add_paragraph(); p.paragraph_format.space_before = Pt(4)
r = p.add_run(SIGNOFF); r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

d.save(OUT_DOCX)
print(f"Wrote {OUT_DOCX}")
